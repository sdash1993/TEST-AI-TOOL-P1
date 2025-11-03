from flask import Flask, request, jsonify, send_from_directory,Response
from flask_cors import CORS
import requests
import json
import os
import sys
import base64
from html.parser import HTMLParser
import traceback
import re
import prompt_v5

from langchain.schema import HumanMessage, AIMessage


from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename
import PyPDF2
from docx import Document
import tempfile
from pathlib import Path

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple

from langchain_google_genai import ChatGoogleGenerativeAI

from prompt_v3 import karate_api_prompt
from prompt_v4 import rest_api_prompt

from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)
CORS(app)

# New Relic API configuration
NEWRELIC_API_URL = "https://api.eu.newrelic.com/graphql"

# Global cache to store exact app names per account
app_cache = {}


class SwaggerAPIParser:
    def __init__(self, base_url: str):
        self.base_url = base_url
        self.headers = {
            'accept': 'application/json,*/*',
            'accept-language': 'en-GB,en-US;q=0.9,en;q=0.8',
            'cache-control': 'no-cache',
            'dnt': '1',
            'pragma': 'no-cache',
            'priority': 'u=1, i',
            'referer': f'{base_url}/swagger-ui/index.html',
            'sec-ch-ua': '"Not)A;Brand";v="8", "Chromium";v="138", "Google Chrome";v="138"',
            'sec-ch-ua-mobile': '?0',
            'sec-ch-ua-platform': '"macOS"',
            'sec-fetch-dest': 'empty',
            'sec-fetch-mode': 'cors',
            'sec-fetch-site': 'same-origin',
            'user-agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/138.0.0.0 Safari/537.36',
            'Cookie': '__cf_bm=b7a5jhNK4aDPD_lMypMMJvixkIQUDf7_XsRKEw_PkmQ-1752057695-1.0.1.1-mA3jlQ5Kb6HDhd89JBaMqcvU8w3ZpDthoHlhxey5UqlRwI1ALbrQcdWpnMqGIIqXl4TrHn3iO6MAhtta5X13Azz5Y8cLMD6CpcpdvIqThEs'
        }
        self.swagger_data = None
        self.components = {}

    def fetch_swagger_data(self) -> bool:
        """Fetch Swagger/OpenAPI specification from the API"""
        try:
            swagger_url = f"{self.base_url}/v3/api-docs"
            response = requests.get(swagger_url, headers=self.headers, timeout=30)
            response.raise_for_status()

            self.swagger_data = response.json()
            self.components = self.swagger_data.get('components', {})
            return True

        except requests.exceptions.RequestException as e:
            print(f"❌ Error fetching Swagger data: {e}")
            return False

    def resolve_schema_ref(self, schema: Dict[str, Any]) -> Dict[str, Any]:
        """Resolve $ref references in schema"""
        if not isinstance(schema, dict):
            return schema

        if '$ref' in schema:
            ref_path = schema['$ref']
            if ref_path.startswith('#/components/schemas/'):
                schema_name = ref_path.split('/')[-1]
                return self.components.get('schemas', {}).get(schema_name, {})

        # Recursively resolve nested schemas
        resolved = {}
        for key, value in schema.items():
            if isinstance(value, dict):
                resolved[key] = self.resolve_schema_ref(value)
            elif isinstance(value, list):
                resolved[key] = [self.resolve_schema_ref(item) if isinstance(item, dict) else item for item in value]
            else:
                resolved[key] = value

        return resolved

    def generate_sample_payload(self, schema: Dict[str, Any]) -> Dict[str, Any]:
        """Generate sample payload from schema"""
        if not isinstance(schema, dict):
            return schema

        # Resolve references first
        resolved_schema = self.resolve_schema_ref(schema)

        sample = {}
        schema_type = resolved_schema.get('type', '')
        properties = resolved_schema.get('properties', {})

        if schema_type == 'object' and properties:
            for prop_name, prop_schema in properties.items():
                sample[prop_name] = self.generate_sample_value(prop_schema)
        elif schema_type == 'array':
            items_schema = resolved_schema.get('items', {})
            sample = [self.generate_sample_value(items_schema)]
        else:
            return self.generate_sample_value(resolved_schema)

        return sample

    def generate_sample_value(self, schema: Dict[str, Any]) -> Any:
        """Generate sample value based on schema type"""
        if not isinstance(schema, dict):
            return schema

        resolved_schema = self.resolve_schema_ref(schema)
        schema_type = resolved_schema.get('type', '')

        # Check for example value first
        if 'example' in resolved_schema:
            return resolved_schema['example']

        # Generate based on type
        if schema_type == 'string':
            enum_values = resolved_schema.get('enum', [])
            if enum_values:
                return enum_values[0]
            return "string"
        elif schema_type == 'integer':
            return 0
        elif schema_type == 'number':
            return 0.0
        elif schema_type == 'boolean':
            return True
        elif schema_type == 'array':
            items_schema = resolved_schema.get('items', {})
            return [self.generate_sample_value(items_schema)]
        elif schema_type == 'object':
            properties = resolved_schema.get('properties', {})
            obj = {}
            for prop_name, prop_schema in properties.items():
                obj[prop_name] = self.generate_sample_value(prop_schema)
            return obj
        else:
            return {}

    def normalize_newrelic_endpoint(self, endpoint_uri: str) -> str:
        """Remove service name from New Relic endpoint for matching"""
        # Remove service names that are common in New Relic but not in Swagger
        service_names_to_remove = [
            '/customs-service',
            '/payment-service',
            '/user-service',
            '/notification-service',
            '/document-service',
            '/tracking-service',
            '/logistics-service',
            '/shipping-service',
            '/cargo-service',
            '/freight-service',
            '/invoice/api',
            '/shipment-service',
            '/runner-order-service',
            '/mdm',
            '/dps-service'

        ]

        normalized_endpoint = endpoint_uri

        for service_name in service_names_to_remove:
            if service_name in normalized_endpoint:
                normalized_endpoint = normalized_endpoint.replace(service_name, '')
                print(f"🔄 NORMALIZED: Removed '{service_name}' from '{endpoint_uri}' -> '{normalized_endpoint}'")
                break

        return normalized_endpoint

    def find_matching_endpoint(self, endpoint_uri: str) -> Tuple[Optional[str], Optional[str]]:
        """Find matching endpoint in Swagger spec with EXACT string matching after removing service name"""
        if not self.swagger_data:
            if not self.fetch_swagger_data():
                return None, None

        paths = self.swagger_data.get('paths', {})

        print(f"🔍 SWAGGER: Original New Relic endpoint: '{endpoint_uri}'")

        # Normalize New Relic endpoint by removing service name
        normalized_nr_endpoint = self.normalize_newrelic_endpoint(endpoint_uri)

        print(f"🔍 SWAGGER: Normalized New Relic endpoint: '{normalized_nr_endpoint}'")
        print(f"🔍 SWAGGER: Available Swagger paths:")
        for i, path in enumerate(paths.keys()):
            print(f"   {i + 1}. '{path}'")

        # EXACT MATCH with normalized endpoint
        if normalized_nr_endpoint in paths:
            print(f"✅ SWAGGER: EXACT MATCH FOUND: '{normalized_nr_endpoint}'")
            # Find the first available method
            for method in ['post', 'put', 'patch', 'get', 'delete']:
                if method in paths[normalized_nr_endpoint]:
                    print(f"✅ SWAGGER: Method '{method.upper()}' found for exact match")
                    return normalized_nr_endpoint, method.upper()

        print(f"❌ SWAGGER: NO EXACT MATCH FOUND")
        print(f"❌ SWAGGER: Looking for exact match: '{normalized_nr_endpoint}'")
        print(f"❌ SWAGGER: Available Swagger paths:")
        for i, path in enumerate(paths.keys()):
            print(f"   {i + 1}. '{path}'")

        return None, None

    def get_endpoint_payload(self, endpoint_path: str, method: str = None) -> Tuple[Optional[Dict], Optional[Dict]]:
        """Get request body and response body for a specific endpoint"""
        if not self.swagger_data:
            if not self.fetch_swagger_data():
                return None, None

        paths = self.swagger_data.get('paths', {})

        # Find the endpoint
        endpoint_data = None
        found_method = None

        if endpoint_path in paths:
            endpoint_methods = paths[endpoint_path]

            if method:
                # Look for specific method
                if method.lower() in endpoint_methods:
                    endpoint_data = endpoint_methods[method.lower()]
                    found_method = method.upper()
            else:
                # Find any available method
                for http_method in ['post', 'put', 'patch', 'get', 'delete']:
                    if http_method in endpoint_methods:
                        endpoint_data = endpoint_methods[http_method]
                        found_method = http_method.upper()
                        break

        if not endpoint_data:
            return None, None

        # Extract request body
        request_body = None
        if 'requestBody' in endpoint_data:
            request_body_data = endpoint_data['requestBody']
            content = request_body_data.get('content', {})

            for media_type, media_data in content.items():
                if media_type == 'application/json':
                    schema = media_data.get('schema', {})
                    request_body = self.generate_sample_payload(schema)
                    break

        # Extract response body
        response_body = None
        if 'responses' in endpoint_data:
            responses = endpoint_data['responses']

            # Try to get 200 response first, then any other success response
            for status_code in ['200', '201', '202', '204']:
                if status_code in responses:
                    response_data = responses[status_code]
                    content = response_data.get('content', {})

                    for media_type, media_data in content.items():
                        if media_type == 'application/json' or media_type == '*/*':
                            schema = media_data.get('schema', {})
                            response_body = self.generate_sample_payload(schema)
                            break

                    if response_body:
                        break

        return request_body, response_body

    def extract_parameters(self, endpoint_path: str, method: str = None) -> Dict[str, Any]:
        """Extract all parameters (path, query, header, body) from endpoint"""
        if not self.swagger_data:
            if not self.fetch_swagger_data():
                return {}

        paths = self.swagger_data.get('paths', {})

        if endpoint_path not in paths:
            return {}

        endpoint_methods = paths[endpoint_path]

        # Find the method
        endpoint_data = None
        found_method = None

        if method:
            if method.lower() in endpoint_methods:
                endpoint_data = endpoint_methods[method.lower()]
                found_method = method.upper()
        else:
            # Find any available method
            for http_method in ['post', 'put', 'patch', 'get', 'delete']:
                if http_method in endpoint_methods:
                    endpoint_data = endpoint_methods[http_method]
                    found_method = http_method.upper()
                    break

        if not endpoint_data:
            return {}

        parameters = {
            'path_parameters': [],
            'query_parameters': [],
            'header_parameters': [],
            'body_parameters': {},
            'method': found_method,
            'endpoint': endpoint_path
        }

        # Extract parameters from the endpoint definition
        endpoint_parameters = endpoint_data.get('parameters', [])

        for param in endpoint_parameters:
            param_info = {
                'name': param.get('name', ''),
                'type': param.get('schema', {}).get('type', 'string'),
                'required': param.get('required', False),
                'description': param.get('description', ''),
                'example': param.get('example', param.get('schema', {}).get('example', '')),
                'enum': param.get('schema', {}).get('enum', []),
                'default': param.get('schema', {}).get('default', '')
            }

            param_location = param.get('in', '')

            if param_location == 'path':
                parameters['path_parameters'].append(param_info)
            elif param_location == 'query':
                parameters['query_parameters'].append(param_info)
            elif param_location == 'header':
                parameters['header_parameters'].append(param_info)

        # Extract body parameters from requestBody
        if 'requestBody' in endpoint_data:
            request_body = endpoint_data['requestBody']
            content = request_body.get('content', {})

            for media_type, media_data in content.items():
                if media_type == 'application/json':
                    schema = media_data.get('schema', {})
                    parameters['body_parameters'] = self.extract_body_parameters(schema)
                    break

        return parameters

    def extract_body_parameters(self, schema: Dict[str, Any]) -> Dict[str, Any]:
        """Extract body parameters from request body schema"""
        resolved_schema = self.resolve_schema_ref(schema)

        if not isinstance(resolved_schema, dict):
            return {}

        body_params = {
            'type': resolved_schema.get('type', 'object'),
            'properties': {},
            'required': resolved_schema.get('required', []),
            'description': resolved_schema.get('description', '')
        }

        properties = resolved_schema.get('properties', {})

        for prop_name, prop_schema in properties.items():
            prop_resolved = self.resolve_schema_ref(prop_schema)

            param_info = {
                'type': prop_resolved.get('type', 'string'),
                'description': prop_resolved.get('description', ''),
                'example': prop_resolved.get('example', ''),
                'enum': prop_resolved.get('enum', []),
                'default': prop_resolved.get('default', ''),
                'required': prop_name in body_params['required']
            }

            # Handle nested objects
            if param_info['type'] == 'object':
                nested_props = prop_resolved.get('properties', {})
                param_info['properties'] = {}
                for nested_name, nested_schema in nested_props.items():
                    nested_resolved = self.resolve_schema_ref(nested_schema)
                    param_info['properties'][nested_name] = {
                        'type': nested_resolved.get('type', 'string'),
                        'description': nested_resolved.get('description', ''),
                        'example': nested_resolved.get('example', ''),
                        'required': nested_name in nested_resolved.get('required', [])
                    }

            # Handle arrays
            elif param_info['type'] == 'array':
                items_schema = prop_resolved.get('items', {})
                items_resolved = self.resolve_schema_ref(items_schema)
                param_info['items'] = {
                    'type': items_resolved.get('type', 'string'),
                    'description': items_resolved.get('description', ''),
                    'example': items_resolved.get('example', '')
                }

            body_params['properties'][prop_name] = param_info

        return body_params

    def generate_curl_command(self, endpoint_path: str, method: str, parameters: Dict[str, Any], base_url: str = None) -> str:
        """Generate cURL command based on endpoint and parameters"""
        if not base_url:
            base_url = self.base_url

        full_url = f"{base_url}{endpoint_path}"

        # Start building cURL command
        curl_parts = [f"curl -X {method}"]

        # Add headers
        headers = []
        if parameters.get('header_parameters'):
            for param in parameters['header_parameters']:
                example_value = param.get('example') or param.get('default') or f"YOUR_{param['name'].upper()}"
                headers.append(f'-H "{param["name"]}: {example_value}"')

        # Add content-type for requests with body
        if method in ['POST', 'PUT', 'PATCH'] and parameters.get('body_parameters', {}).get('properties'):
            headers.append('-H "Content-Type: application/json"')

        # Add headers to curl command
        for header in headers:
            curl_parts.append(f"  {header}")

        # Handle path parameters (replace in URL)
        if parameters.get('path_parameters'):
            for param in parameters['path_parameters']:
                example_value = param.get('example') or param.get('default') or f"{{{param['name']}}}"
                full_url = full_url.replace(f"{{{param['name']}}}", str(example_value))

        # Handle query parameters
        query_params = []
        if parameters.get('query_parameters'):
            for param in parameters['query_parameters']:
                example_value = param.get('example') or param.get('default') or f"YOUR_{param['name'].upper()}"
                query_params.append(f"{param['name']}={example_value}")

        if query_params:
            full_url += "?" + "&".join(query_params)

        # Add URL to curl command
        curl_parts.append(f'  "{full_url}"')

        # Handle body parameters
        if method in ['POST', 'PUT', 'PATCH'] and parameters.get('body_parameters', {}).get('properties'):
            body_example = {}
            for prop_name, prop_info in parameters['body_parameters']['properties'].items():
                example_value = prop_info.get('example') or prop_info.get('default')
                if not example_value:
                    if prop_info['type'] == 'string':
                        example_value = f"YOUR_{prop_name.upper()}"
                    elif prop_info['type'] == 'integer':
                        example_value = 0
                    elif prop_info['type'] == 'number':
                        example_value = 0.0
                    elif prop_info['type'] == 'boolean':
                        example_value = True
                    elif prop_info['type'] == 'array':
                        example_value = []
                    elif prop_info['type'] == 'object':
                        example_value = {}
                    else:
                        example_value = f"YOUR_{prop_name.upper()}"

                body_example[prop_name] = example_value

            body_json = json.dumps(body_example, indent=2)
            curl_parts.append(f"  -d '{body_json}'")

        return " \\\n".join(curl_parts)


def get_available_accounts(api_key):
    """Get all accounts accessible with the given API key"""
    headers = {
        'Content-Type': 'application/json',
        'API-Key': api_key
    }

    graphql_query = {
        "query": """
        {
            actor {
                accounts {
                    id
                    name
                }
            }
        }
        """
    }

    try:
        response = requests.post(NEWRELIC_API_URL,
                                 headers=headers,
                                 json=graphql_query,
                                 timeout=30)
        response.raise_for_status()
        return response.json()
    except requests.exceptions.RequestException as e:
        raise Exception(f"API request failed: {str(e)}")


def sanitize_app_name_for_nrql(app_name):
    """
    Sanitize app name for NRQL queries while preserving exact value
    ONLY escapes single quotes for SQL safety
    """
    if not app_name:
        raise ValueError("App name cannot be empty")

    # Only escape single quotes for NRQL (SQL-style escaping)
    sanitized = app_name.replace("'", "''")

    print(f"✅ SANITIZED APP NAME: '{app_name}' -> '{sanitized}'")
    return sanitized


def query_newrelic(account_id, api_key, nrql_query):
    """Execute NRQL query against New Relic API"""
    headers = {
        'Content-Type': 'application/json',
        'API-Key': api_key
    }

    graphql_query = {
        "query": """
        query($accountId: Int!, $nrql: Nrql!) {
            actor {
                account(id: $accountId) {
                    nrql(query: $nrql) {
                        results
                    }
                }
            }
        }
        """,
        "variables": {
            "accountId": int(account_id),
            "nrql": nrql_query
        }
    }

    try:
        response = requests.post(NEWRELIC_API_URL,
                                 headers=headers,
                                 json=graphql_query,
                                 timeout=60)
        response.raise_for_status()
        result = response.json()

        if 'errors' in result:
            error_msg = result['errors'][0]['message']
            if 'timeout' in error_msg.lower():
                raise Exception(f"Query timeout - try a shorter time range or simpler query")
            raise Exception(f"GraphQL Error: {error_msg}")

        return result
    except requests.exceptions.RequestException as e:
        raise Exception(f"API request failed: {str(e)}")


def safe_numeric_value(value, default=0):
    """Safely convert value to numeric, handling various data types"""
    if value is None:
        return default

    # If it's already a number, return it
    if isinstance(value, (int, float)):
        return value

    # If it's a string, try to convert
    if isinstance(value, str):
        try:
            return float(value)
        except ValueError:
            return default

    # If it's a dict (sometimes New Relic returns nested data), try to extract numeric value
    if isinstance(value, dict):
        # Common patterns in New Relic responses
        for key in ['value', 'count', 'average', 'sum']:
            if key in value:
                return safe_numeric_value(value[key], default)
        return default

    # If it's a list, try to get the first numeric value
    if isinstance(value, list) and value:
        return safe_numeric_value(value[0], default)

    return default


def normalize_api_data(raw_data):
    """Normalize API data from New Relic to consistent format"""
    normalized = []

    for item in raw_data:
        if not isinstance(item, dict):
            continue

        # Extract API endpoint name
        api_name = item.get('request.uri', item.get('facet', 'Unknown'))

        # Skip if no valid API name
        if not api_name or api_name == 'Unknown':
            continue

        # Normalize the data structure
        normalized_item = {
            'api_endpoint': api_name,
            'total_requests': safe_numeric_value(item.get('requests', item.get('count', 0))),
            'avg_response_time': safe_numeric_value(item.get('avg_response_time', item.get('average', 0))),
            'min_response_time': safe_numeric_value(item.get('min_response_time', item.get('min', 0))),
            'max_response_time': safe_numeric_value(item.get('max_response_time', item.get('max', 0))),
            'error_count': safe_numeric_value(item.get('error_count', 0)),
            'success_rate': safe_numeric_value(item.get('success_rate', 100)),
            'first_seen': item.get('first_seen', 'Unknown'),
            'last_seen': item.get('last_seen', 'Unknown'),
            # Keep original fields for compatibility
            'request.uri': api_name,
            'count': safe_numeric_value(item.get('count', item.get('requests', 0))),
            'average': safe_numeric_value(item.get('average', item.get('avg_response_time', 0))),
            'max': safe_numeric_value(item.get('max', item.get('max_response_time', 0))),
            'min': safe_numeric_value(item.get('min', item.get('min_response_time', 0)))
        }

        normalized.append(normalized_item)

    return normalized


def get_available_apps(account_id, api_key):
    """Get all available apps - returns exact app names as they appear in New Relic"""
    print(f"✅ GETTING APPS FOR ACCOUNT: {account_id}")

    queries_to_try = [
        {
            'name': 'Last 30 days',
            'query': "SELECT uniques(appName) FROM Transaction SINCE 30 days ago LIMIT 100",
            'primary': True
        },
        {
            'name': 'Last 7 days',
            'query': "SELECT uniques(appName) FROM Transaction SINCE 7 days ago LIMIT 100",
            'primary': False
        },
        {
            'name': 'Last 90 days',
            'query': "SELECT uniques(appName) FROM Transaction SINCE 90 days ago LIMIT 100",
            'primary': False
        }
    ]

    all_apps = set()
    successful_query = None

    for query_info in queries_to_try:
        try:
            print(f"Trying query: {query_info['name']} - {query_info['query']}")
            result = query_newrelic(account_id, api_key, query_info['query'])

            if 'data' in result and result['data']['actor']['account']['nrql']:
                apps_data = result['data']['actor']['account']['nrql']['results']

                if apps_data:
                    for item in apps_data:
                        if 'uniques.appName' in item:
                            app_names = item['uniques.appName']
                            if isinstance(app_names, list):
                                for app_name in app_names:
                                    if app_name and app_name.strip():
                                        exact_app_name = app_name.strip()
                                        all_apps.add(exact_app_name)
                                        print(f"✅ FOUND EXACT APP: '{exact_app_name}'")
                        elif 'appName' in item and item['appName']:
                            exact_app_name = item['appName'].strip()
                            if exact_app_name:
                                all_apps.add(exact_app_name)
                                print(f"✅ FOUND EXACT APP: '{exact_app_name}'")

                    if len(all_apps) > 0:
                        successful_query = query_info
                        print(f"✅ SUCCESS: Found {len(all_apps)} apps with {query_info['name']} query")
                        break

        except Exception as e:
            print(f"❌ Query {query_info['name']} failed: {str(e)}")
            continue

    apps_list = sorted(list(all_apps))
    cache_key = f"{account_id}_{api_key[:10]}"
    app_cache[cache_key] = apps_list

    print(f"✅ FINAL EXACT APPS LIST: {apps_list}")
    return apps_list, successful_query

@app.route('/ai')
def serve_landing_page():
    """Serve the landing page"""
    if os.path.exists('UI/llmapidiscovery.html'):
        return send_from_directory('UI', 'llmapidiscovery.html')
    else:
        return "Landing page not found", 404


@app.route('/ai/llmapidiscovery')
def serve_llm_discovery():
    """Serve the LLM API Discovery page"""
    if os.path.exists('UI/index.html'):
        return send_from_directory('UI', 'index.html')
    else:
        return "LLM API Discovery page not found", 404

def validate_app_name(app_name, account_id, api_key):
    """Validate that the app name exists in the cached apps"""
    cache_key = f"{account_id}_{api_key[:10]}"

    if cache_key not in app_cache:
        print(f"❌ NO CACHED APPS FOR ACCOUNT {account_id}")
        return False, "No cached apps found. Please get apps first."

    cached_apps = app_cache[cache_key]

    if app_name not in cached_apps:
        print(f"❌ APP '{app_name}' NOT FOUND IN CACHED APPS: {cached_apps}")
        return False, f"App '{app_name}' not found in available applications"

    print(f"✅ APP '{app_name}' VALIDATED SUCCESSFULLY")
    return True, "App validated"


@app.route('/ai/get-accounts', methods=['POST'])
def get_accounts():
    """Get available New Relic accounts"""
    try:
        data = request.json
        api_key = data.get('apiKey')

        if not api_key:
            return jsonify({
                "status": "error",
                "message": "API key is required"
            }), 400

        print(f"Getting accounts with API key: {api_key[:10]}...")
        result = get_available_accounts(api_key)

        if 'data' in result and result['data']['actor']['accounts']:
            accounts = result['data']['actor']['accounts']
            return jsonify({
                "status": "success",
                "accounts": accounts,
                "message": f"Found {len(accounts)} accessible accounts"
            })
        else:
            return jsonify({
                "status": "error",
                "message": "No accounts found or access denied"
            }), 403

    except Exception as e:
        print(f"Error in get_accounts: {str(e)}")
        return jsonify({
            "status": "error",
            "message": str(e)
        }), 500


@app.route('/ai/get-apps', methods=['POST'])
def get_apps():
    """Get available applications - returns exact names as they appear in New Relic"""
    try:
        data = request.json
        account_id = data.get('accountId')
        api_key = data.get('apiKey')

        if not all([account_id, api_key]):
            return jsonify({
                "status": "error",
                "message": "Account ID and API key are required"
            }), 400

        apps, successful_query = get_available_apps(account_id, api_key)

        return jsonify({
            "status": "success",
            "apps": apps,
            "successful_query": successful_query['name'] if successful_query else None,
            "message": f"Found {len(apps)} applications" + (
                f" using {successful_query['name']} query" if successful_query else " - no apps found")
        })

    except Exception as e:
        print(f"Error in get_apps: {str(e)}")
        return jsonify({
            "status": "error",
            "message": str(e)
        }), 500


@app.route('/ai/test-connection', methods=['POST'])
def test_connection():
    """Test connection using EXACT app name from frontend"""
    try:
        data = request.json
        account_id = data.get('accountId')
        api_key = data.get('apiKey')
        app_name = data.get('appName')

        print(f"✅ TESTING CONNECTION WITH EXACT APP NAME: '{app_name}'")

        if not all([account_id, api_key, app_name]):
            return jsonify({
                "status": "error",
                "message": "Account ID, API key, and app name are required"
            }), 400

        is_valid, validation_message = validate_app_name(app_name, account_id, api_key)
        if not is_valid:
            return jsonify({
                "status": "error",
                "message": validation_message
            }), 400

        sanitized_app_name = sanitize_app_name_for_nrql(app_name)
        test_query = f"SELECT count(*) FROM Transaction WHERE appName = '{sanitized_app_name}' SINCE 1 day ago"
        print(f"✅ TEST QUERY: {test_query}")

        result = query_newrelic(account_id, api_key, test_query)

        if 'data' in result and result['data']['actor']['account']['nrql']:
            test_data = result['data']['actor']['account']['nrql']['results']
            if test_data and len(test_data) > 0:
                count = safe_numeric_value(test_data[0].get('count', 0))
                return jsonify({
                    "status": "success",
                    "message": "Connection successful",
                    "app_name": app_name,
                    "app_transactions": count,
                    "app_status": "found" if count > 0 else "no_transactions"
                })
            else:
                return jsonify({
                    "status": "error",
                    "message": f"No data returned for app '{app_name}'"
                }), 400
        else:
            return jsonify({
                "status": "error",
                "message": "Invalid response from New Relic API"
            }), 400

    except Exception as e:
        print(f"Error in test_connection: {str(e)}")
        return jsonify({
            "status": "error",
            "message": str(e)
        }), 500


@app.route('/ai/get-endpoint-details', methods=['POST'])
def get_endpoint_details():
    """Get detailed information for a specific endpoint including enhanced Swagger payloads and parameters"""
    try:
        data = request.json
        account_id = data.get('accountId')
        api_key = data.get('apiKey')
        app_name = data.get('appName')
        endpoint_uri = data.get('endpointUri')
        time_range = data.get('timeRange', 90)
        swagger_base_url = data.get('swaggerBaseUrl')  # Optional Swagger URL

        print(f"✅ GETTING ENDPOINT DETAILS FOR: '{endpoint_uri}'")
        print(f"✅ SWAGGER BASE URL: '{swagger_base_url}'")

        if not all([account_id, api_key, app_name, endpoint_uri]):
            return jsonify({
                "status": "error",
                "message": "Account ID, API key, app name, and endpoint URI are required"
            }), 400

        is_valid, validation_message = validate_app_name(app_name, account_id, api_key)
        if not is_valid:
            return jsonify({
                "status": "error",
                "message": validation_message
            }), 400

        sanitized_app_name = sanitize_app_name_for_nrql(app_name)
        sanitized_endpoint = endpoint_uri.replace("'", "''")  # Escape single quotes for NRQL

        # Query to get detailed endpoint information from New Relic
        detail_query = f"""
        SELECT
            earliest(request.uri) as endpoint,
            earliest(request.method) as httpMethod,
            earliest(request.headers.contentType) as contentType,
            earliest(request.headers.accept) as accept,
            earliest(request.body) as requestBody,
            earliest(request.payload) as requestPayload
        FROM Transaction
        WHERE appName = '{sanitized_app_name}'
        AND request.uri = '{sanitized_endpoint}'
        SINCE {time_range} days ago
        """

        print(f"✅ ENDPOINT DETAIL QUERY: {detail_query}")

        result = query_newrelic(account_id, api_key, detail_query)

        endpoint_details = {
            'endpoint': endpoint_uri,
            'httpMethod': 'Unknown',
            'contentType': 'Unknown',
            'accept': 'Unknown',
            'requestBody': 'No body data',
            'requestPayload': 'No payload data',
            'swagger_request_payload': None,
            'swagger_response_payload': None,
            'swagger_matched': False,
            'swagger_matched_path': None,
            'swagger_matched_method': None,
            'swagger_error': None,
            'swagger_parameters': None,
            'swagger_curl_command': None
        }

        # Get New Relic data
        if 'data' in result and result['data']['actor']['account']['nrql']:
            detail_data = result['data']['actor']['account']['nrql']['results']

            if detail_data and len(detail_data) > 0:
                nr_details = detail_data[0]
                endpoint_details.update({
                    'endpoint': nr_details.get('endpoint', endpoint_uri),
                    'httpMethod': nr_details.get('httpMethod', 'Unknown'),
                    'contentType': nr_details.get('contentType', 'Unknown'),
                    'accept': nr_details.get('accept', 'Unknown'),
                    'requestBody': nr_details.get('requestBody', 'No body data'),
                    'requestPayload': nr_details.get('requestPayload', 'No payload data')
                })

        # Try to get enhanced Swagger data if base URL is provided
        if swagger_base_url and swagger_base_url.strip():
            try:
                print(f"✅ ATTEMPTING ENHANCED SWAGGER INTEGRATION FOR: {endpoint_uri}")
                swagger_parser = SwaggerAPIParser(swagger_base_url.strip())

                # Try to find matching endpoint
                swagger_path, swagger_method = swagger_parser.find_matching_endpoint(endpoint_uri)

                if swagger_path and swagger_method:
                    print(f"✅ FOUND SWAGGER MATCH: {swagger_method} {swagger_path}")

                    # Get payloads from Swagger
                    request_payload, response_payload = swagger_parser.get_endpoint_payload(swagger_path, swagger_method)

                    # Get enhanced parameters
                    parameters = swagger_parser.extract_parameters(swagger_path, swagger_method)

                    # Generate cURL command
                    curl_command = swagger_parser.generate_curl_command(swagger_path, swagger_method, parameters, swagger_base_url)

                    endpoint_details.update({
                        'swagger_request_payload': request_payload,
                        'swagger_response_payload': response_payload,
                        'swagger_matched': True,
                        'swagger_matched_path': swagger_path,
                        'swagger_matched_method': swagger_method,
                        'swagger_parameters': parameters,
                        'swagger_curl_command': curl_command
                    })

                    print(f"✅ ENHANCED SWAGGER DATA EXTRACTED SUCCESSFULLY")
                    print(f"✅ REQUEST PAYLOAD: {request_payload is not None}")
                    print(f"✅ RESPONSE PAYLOAD: {response_payload is not None}")
                    print(f"✅ PARAMETERS: {parameters is not None}")
                    print(f"✅ CURL COMMAND: {curl_command is not None}")
                else:
                    print(f"❌ NO SWAGGER MATCH FOUND FOR: {endpoint_uri}")
                    endpoint_details['swagger_error'] = f"No matching endpoint found in Swagger for {endpoint_uri}"

            except Exception as swagger_error:
                print(f"❌ SWAGGER ERROR: {swagger_error}")
                endpoint_details['swagger_error'] = f"Swagger integration error: {str(swagger_error)}"
        else:
            print(f"ℹ️ NO SWAGGER URL PROVIDED - SKIPPING SWAGGER INTEGRATION")

        return jsonify({
            "status": "success",
            "endpoint_details": endpoint_details,
            "query_used": detail_query,
            "time_range": time_range,
            "swagger_integration": swagger_base_url is not None and swagger_base_url.strip() != ""
        })

    except Exception as e:
        print(f"Error in get_endpoint_details: {str(e)}")
        return jsonify({
            "status": "error",
            "message": str(e)
        }), 500


@app.route('/ai/analyze-data', methods=['POST'])
def analyze_data():
    """Analyze data to get all APIs from last 30 days"""
    try:
        data = request.json
        account_id = data.get('accountId')
        api_key = data.get('apiKey')
        app_name = data.get('appName')
        time_range = data.get('timeRange', 30)

        print(f"✅ ANALYZING API DATA FOR APP: '{app_name}'")
        print(f"✅ TIME RANGE: {time_range} days")

        if not all([account_id, api_key, app_name]):
            return jsonify({
                "status": "error",
                "message": "Account ID, API key, and app name are required"
            }), 400

        is_valid, validation_message = validate_app_name(app_name, account_id, api_key)
        if not is_valid:
            return jsonify({
                "status": "error",
                "message": validation_message
            }), 400

        sanitized_app_name = sanitize_app_name_for_nrql(app_name)

        # Queries focused on API discovery and basic metrics
        queries = {
            'all_apis': f"SELECT count(*) as requests, average(duration) as avg_response_time, min(duration) as min_response_time, max(duration) as max_response_time FROM Transaction WHERE appName = '{sanitized_app_name}' FACET request.uri SINCE {time_range} days ago LIMIT 200",
            'summary': f"SELECT count(*) as total_requests, average(duration) as avg_response_time, min(duration) as min_response_time, max(duration) as max_response_time FROM Transaction WHERE appName = '{sanitized_app_name}' SINCE {time_range} days ago",
            'traffic_patterns': f"SELECT count(*) as requests FROM Transaction WHERE appName = '{sanitized_app_name}' FACET hourOf(timestamp) SINCE 7 days ago LIMIT 168",
            'daily_volume': f"SELECT count(*) as requests FROM Transaction WHERE appName = '{sanitized_app_name}' FACET dateOf(timestamp) SINCE {time_range} days ago LIMIT 30",
            'error_analysis': f"SELECT count(*) as errors FROM TransactionError WHERE appName = '{sanitized_app_name}' FACET request.uri SINCE {time_range} days ago LIMIT 100"
        }

        print("✅ EXECUTING API DISCOVERY QUERIES:")
        for query_name, query in queries.items():
            print(f"  {query_name}: {query}")

        # Execute all queries
        results = {}
        for query_name, query in queries.items():
            try:
                result = query_newrelic(account_id, api_key, query)
                raw_results = result.get('data', {}).get('actor', {}).get('account', {}).get('nrql', {}).get('results',
                                                                                                             [])
                results[query_name] = raw_results
                print(f"✅ Query '{query_name}' returned {len(raw_results)} results")

                # Debug: Print first result to see data structure
                if raw_results and len(raw_results) > 0:
                    print(f"🔍 First result for '{query_name}': {raw_results[0]}")

            except Exception as e:
                print(f"❌ Query '{query_name}' failed: {e}")
                results[query_name] = []

        # Process results
        all_apis = normalize_api_data(results['all_apis'])
        summary_results = results['summary']
        traffic_patterns = results['traffic_patterns']
        daily_volume = results['daily_volume']
        error_analysis = results['error_analysis']

        # Calculate summary metrics
        summary_metrics = summary_results[0] if summary_results else {}

        # Process error data to add error counts to APIs
        error_map = {}
        for error_item in error_analysis:
            uri = error_item.get('request.uri', error_item.get('facet', ''))
            if uri:
                error_map[uri] = safe_numeric_value(error_item.get('errors', error_item.get('count', 0)))

        # Add error counts to API data
        for api in all_apis:
            api_uri = api['api_endpoint']
            api['error_count'] = error_map.get(api_uri, 0)
            # Calculate success rate
            total_requests = api['total_requests']
            if total_requests > 0:
                api['success_rate'] = round(((total_requests - api['error_count']) / total_requests) * 100, 2)
            else:
                api['success_rate'] = 100.0

        # Sort APIs by total requests (most popular first)
        all_apis = sorted(all_apis, key=lambda x: x['total_requests'], reverse=True)

        # Create categorized lists
        high_traffic_apis = [api for api in all_apis if api['total_requests'] > 1000][:20]
        slow_apis = [api for api in all_apis if api['avg_response_time'] > 1000][:20]
        error_prone_apis = [api for api in all_apis if api['error_count'] > 10][:20]
        all_unique_apis = all_apis  # All APIs found

        # Calculate statistics
        api_stats = {
            'total_apis_found': len(all_apis),
            'high_traffic_apis': len(high_traffic_apis),
            'slow_apis': len(slow_apis),
            'error_prone_apis': len(error_prone_apis),
            'total_requests': sum(api['total_requests'] for api in all_apis),
            'avg_response_time_all': round(sum(api['avg_response_time'] for api in all_apis) / len(all_apis),
                                           2) if all_apis else 0,
            'total_errors': sum(api['error_count'] for api in all_apis)
        }

        return jsonify({
            "status": "success",
            "app_name": app_name,
            "time_range_used": time_range,
            "summary": {
                "total_apis": len(all_apis),
                "total_requests": int(safe_numeric_value(summary_metrics.get('total_requests', 0))),
                "avg_response_time": round(safe_numeric_value(summary_metrics.get('avg_response_time', 0)), 2),
                "min_response_time": round(safe_numeric_value(summary_metrics.get('min_response_time', 0)), 2),
                "max_response_time": round(safe_numeric_value(summary_metrics.get('max_response_time', 0)), 2),
                "analysis_period": f"Last {time_range} days"
            },
            "all_apis": all_unique_apis,
            "categorized_apis": {
                "high_traffic": high_traffic_apis,
                "slow_response": slow_apis,
                "error_prone": error_prone_apis
            },
            "api_statistics": api_stats,
            "traffic_patterns": traffic_patterns,
            "daily_volume": daily_volume,
            "queries_executed": list(queries.keys()),
            # Add connection data for endpoint details
            "connection_data": {
                "account_id": account_id,
                "api_key": api_key,
                "app_name": app_name,
                "time_range": time_range
            }
        })

    except Exception as e:
        print(f"❌ Analysis error: {str(e)}")
        return jsonify({
            "status": "error",
            "message": str(e)
        }), 500


@app.route('/ai/export-data', methods=['POST'])
def export_data():
    """Export analyzed API data"""
    try:
        data = request.json
        export_format = data.get('format', 'json')
        notes = data.get('notes', '')
        analysis_data = data.get('analysisData', {})

        export_package = {
            "timestamp": datetime.now().isoformat(),
            "format": export_format,
            "notes": notes,
            "analysis": analysis_data,
            "metadata": {
                "generated_by": "API Automation Suite",
                "version": "1.0.0",
                "focus": "API Discovery and Enhanced Parameter Extraction"
            }
        }

        filename = f"api_analysis_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"

        return jsonify({
            "status": "success",
            "message": "Export package generated successfully",
            "filename": filename,
            "data": export_package
        })

    except Exception as e:
        return jsonify({
            "status": "error",
            "message": str(e)
        }), 500


@app.route('/ai/validate-app', methods=['POST'])
def validate_app():
    """Validate that the selected app name is available"""
    try:
        data = request.json
        account_id = data.get('accountId')
        api_key = data.get('apiKey')
        app_name = data.get('appName')

        is_valid, message = validate_app_name(app_name, account_id, api_key)

        cache_key = f"{account_id}_{api_key[:10]}"
        available_apps_count = len(app_cache.get(cache_key, []))

        return jsonify({
            "status": "success",
            "is_valid": is_valid,
            "app_name": app_name,
            "message": message,
            "available_apps_count": available_apps_count
        })

    except Exception as e:
        return jsonify({
            "status": "error",
            "message": str(e)
        }), 500


@app.errorhandler(404)
def not_found(error):
    return jsonify({"error": "Endpoint not found"}), 404


@app.errorhandler(500)
def internal_error(error):
    return jsonify({"error": "Internal server error"}), 500

AVAILABLE_PROMPTS = {
    "karate_api": {
        "name": "Generate Karate API Test",
        "description": "Generate Karate Test scripts for API testing",
        "prompt": karate_api_prompt,
        "category": "Karate Test"
     }
    ,
    "rest_assured": {
        "name": "Generate Rest API Test",
        "description": "Generate Rest Assured Test scripts for API testing",
        "prompt": rest_api_prompt,
        "category": "Rest APi Test"
    }
}


def contact_llm_with_prompt(endpoint, method_type, payload=None, prompt_type="karate_api"):
    """Contact LLM with specified prompt type"""
    if prompt_type not in AVAILABLE_PROMPTS:
        raise ValueError(f"Unknown prompt type: {prompt_type}")

    prompt_config = AVAILABLE_PROMPTS[prompt_type]
    selected_prompt = prompt_config["prompt"]

    # Handle module-level prompts (like prompt_v4)
    if hasattr(selected_prompt, '__name__') and selected_prompt.__name__.startswith('prompt_v'):
        # If it's a module, look for a main prompt variable
        # Common patterns: main_prompt, default_prompt, or a prompt with the module name
        if hasattr(selected_prompt, 'main_prompt'):
            selected_prompt = selected_prompt.main_prompt
        elif hasattr(selected_prompt, 'default_prompt'):
            selected_prompt = selected_prompt.default_prompt
        elif hasattr(selected_prompt, prompt_type + '_prompt'):
            selected_prompt = getattr(selected_prompt, prompt_type + '_prompt')
        else:
            # Look for any prompt-like variable in the module
            prompt_vars = [attr for attr in dir(selected_prompt) if
                           'prompt' in attr.lower() and not attr.startswith('_')]
            if prompt_vars:
                selected_prompt = getattr(selected_prompt, prompt_vars[0])
            else:
                raise ValueError(f"No prompt found in module {selected_prompt.__name__}")

    llm = ChatGoogleGenerativeAI(model="models/gemini-2.0-flash", api_key=os.getenv('GOOGLE_API_KEY'))
    chain = selected_prompt | llm
    response = chain.invoke({
        "endpoint": endpoint,
        "method_type": method_type,
        "payload": payload
    })
    return response.content


@app.route('/ai/get-available-prompts', methods=['GET'])
def get_available_prompts():
    """Get list of available prompts for UI selection"""
    try:
        prompts_info = []
        for prompt_id, config in AVAILABLE_PROMPTS.items():
            prompts_info.append({
                "id": prompt_id,
                "name": config["name"],
                "description": config["description"],
                "category": config["category"]
            })

        return jsonify({
            "status": "success",
            "prompts": prompts_info,
            "total_prompts": len(prompts_info)
        }), 200

    except Exception as e:
        return jsonify({
            "status": "error",
            "message": f"Failed to get available prompts: {str(e)}"
        }), 500


@app.route('/ai/contact-llm', methods=['POST'])
def contact_llm_api():
    """API endpoint to contact LLM with endpoint, method_type, optional payload, and prompt type in request body"""
    try:
        # Extract parameters from request body
        data = request.json

        if not data:
            return jsonify({
                "status": "error",
                "message": "Request body is required"
            }), 400

        endpoint = data.get('endpoint')
        method_type = data.get('method_type')
        payload = data.get('payload', None)
        prompt_type = data.get('prompt_type', 'karate_api')  # Default to karate_api

        # Validate required fields
        if not endpoint:
            return "endpoint is required", 400
        if not method_type:
            return "method_type is required", 400

        # Validate prompt type
        if prompt_type not in AVAILABLE_PROMPTS:
            available_types = list(AVAILABLE_PROMPTS.keys())
            return f"Invalid prompt_type '{prompt_type}'. Available types: {available_types}", 400

        prompt_config = AVAILABLE_PROMPTS[prompt_type]

        print(f"✅ CONTACTING LLM:")
        print(f"  Endpoint: {endpoint}")
        print(f"  Method Type: {method_type}")
        print(f"  Payload: {'Provided' if payload is not None else 'None'}")
        print(f"  Prompt Type: {prompt_type} ({prompt_config['name']})")

        # Call the LLM with specified prompt
        response = contact_llm_with_prompt(
            endpoint=endpoint,
            method_type=method_type,
            payload=payload,
            prompt_type=prompt_type
        )

        print(f"✅ LLM RESPONSE RECEIVED")
        print(response)

        # Return JSON response with the LLM response and metadata
        return response, 200

    except Exception as e:
        print(f"❌ ERROR in contact_llm_api: {str(e)}")
        return jsonify({
            "status": "error",
            "message": f"Failed to contact LLM: {str(e)}"
        }), 500

@app.route('/ai/contact-testCase-llm', methods=['POST'])
def contact_llm_testCase_api():
    """API endpoint to contact LLM with Epic data, User Story data, and PRD data"""
    try:
        # Extract parameters from request body
        data = request.json

        if not data:
            return jsonify({
                "status": "error",
                "message": "Request body is required"
            }), 400

        # Extract optional parameters
        epic_data = data.get('epic_data', '')
        us_data = data.get('us_data', '')
        prd_data = data.get('prd_data', '')

        # At least one parameter should be provided
        if not epic_data and not us_data and not prd_data:
            return jsonify({
                "status": "error",
                "message": "At least one of epic_data, us_data, or prd_data must be provided"
            }), 400

        print(f"✅ CONTACTING LLM:")
        print(f"  Epic Data: {'Provided' if epic_data else 'None'} - Length: {len(epic_data) if epic_data else 0}")
        print(f"  US Data: {'Provided' if us_data else 'None'} - Length: {len(us_data) if us_data else 0}")
        print(f"  PRD Data: {'Provided' if prd_data else 'None'} - Length: {len(prd_data) if prd_data else 0}")

        # Call the LLM with the provided data
        response = contact_llm_with_testCase_prompt(
            epic_data=epic_data,
            us_data=us_data,
            prd_data=prd_data
        )

        print(f"✅ LLM RESPONSE RECEIVED")
        print(response)

        # Return JSON response with the LLM response
        return response, 200

    except Exception as e:
        print(f"❌ ERROR in contact_llm_api: {str(e)}")
        return jsonify({
            "status": "error",
            "message": f"Failed to contact LLM: {str(e)}"
        }), 500


def contact_llm_with_testCase_prompt(epic_data='', us_data='', prd_data=''):
    llm = ChatGoogleGenerativeAI(model="models/gemini-2.0-flash", api_key=os.getenv('GOOGLE_API_KEY'))
    chain = prompt_v5.test_case_prompt | llm
    response = chain.invoke({
        "epic_data": epic_data,
        "us_data": us_data,
        "prd_data": prd_data
    })
    return response.content


class HTMLStripper(HTMLParser):
    """Custom HTML parser to strip tags and convert to plain text"""

    def __init__(self):
        super().__init__()
        self.reset()
        self.strict = False
        self.convert_charrefs = True
        self.text = []

    def handle_data(self, data):
        self.text.append(data)

    def get_data(self):
        return ''.join(self.text)


def clean_html_content(html_text):
    """Clean HTML content and format it nicely"""
    if not html_text:
        return ""

    # Replace common HTML patterns with newlines BEFORE stripping
    html_text = re.sub(r'<br\s*/?>', '\n', html_text, flags=re.IGNORECASE)
    html_text = re.sub(r'</div>\s*<div>', '\n', html_text, flags=re.IGNORECASE)
    html_text = re.sub(r'</p>\s*<p>', '\n\n', html_text, flags=re.IGNORECASE)

    # Strip HTML tags
    s = HTMLStripper()
    s.feed(html_text)
    text = s.get_data()

    # Clean up excessive whitespace
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    lines = text.split('\n')
    lines = [line.strip() for line in lines]
    text = '\n'.join(lines)

    # Remove newlines - replace with space
    text = text.replace('\n', ' ')

    # Remove ALL quotes (both single and double) if you don't need them
    # text = text.replace('"', '').replace("'", '')

    # OR replace quotes with a different character
    text = text.replace('"', "'")

    # Remove backslashes
    text = text.replace('\\', '')

    # Clean up multiple spaces
    text = re.sub(r'\s+', ' ', text)

    # Remove ALL special characters except comma and period
    # This keeps only letters (a-z, A-Z), numbers (0-9), spaces, commas, and periods
    text = re.sub(r'[^a-zA-Z0-9\s,.]', '', text)

    # Clean up multiple spaces
    text = re.sub(r'\s+', ' ', text)

    # Clean up multiple periods or commas
    text = re.sub(r'\.+', '.', text)
    text = re.sub(r',+', ',', text)

    # Clean up spaces before/after punctuation
    text = re.sub(r'\s+([,.])', r'\1', text)
    text = re.sub(r'([,.])\s+', r'\1 ', text)

    return text.strip()

@app.route('/ai/get-api-info', methods=['POST'])
def get_api_info():
    """Simplified API info from Swagger"""
    try:
        data = request.json
        swagger_url = data.get('swaggerUrl')
        api_endpoint = data.get('apiEndpoint')

        if not swagger_url or not api_endpoint:
            return jsonify({
                "status": "error",
                "message": "Both swaggerUrl and apiEndpoint are required"
            }), 400

        swagger_parser = SwaggerAPIParser(swagger_url)

        if not swagger_parser.fetch_swagger_data():
            return jsonify({
                "status": "error",
                "message": "Failed to fetch Swagger documentation"
            }), 500

        matched_path, method = swagger_parser.find_matching_endpoint(api_endpoint)

        if not matched_path or not method:
            return jsonify({
                "status": "error",
                "message": f"Endpoint not found: {api_endpoint}"
            }), 404

        parameters = swagger_parser.extract_parameters(matched_path, method)
        request_payload, response_payload = swagger_parser.get_endpoint_payload(matched_path, method)

        return jsonify({
            "method": method,
            "endpoint": matched_path,
            "parameters": parameters,
            "request": request_payload,
            "response": response_payload
        })

    except Exception as e:
        return jsonify({
            "status": "error",
            "message": str(e)
        }), 500


@app.route('/ai/api/workitems', methods=['POST'])
def get_work_items():
    """
    Get work items from Azure DevOps using the batch API
    """
    try:
        # Get JSON payload
        data = request.get_json()
        print(f"Received data: {data}")  # Debug log

        # Validate payload
        if not data:
            return jsonify({
                "error": "No JSON payload provided",
                "debug": "Request body is empty or not valid JSON"
            }), 400

        # Check for token
        if 'token' not in data:
            return jsonify({
                "error": "Missing required field",
                "message": "'token' is required",
                "received_fields": list(data.keys())
            }), 400

        token = data['token']

        # Handle both single and multiple work items
        work_item_ids = []

        if 'work_item_id' in data:
            work_item_ids = [data['work_item_id']]
        elif 'work_item_ids' in data:
            work_item_ids = data['work_item_ids']
            if not isinstance(work_item_ids, list):
                return jsonify({
                    "error": "Invalid format",
                    "message": "'work_item_ids' must be a list",
                    "received_type": type(work_item_ids).__name__
                }), 400
        else:
            return jsonify({
                "error": "Missing required field",
                "message": "Either 'work_item_id' or 'work_item_ids' is required",
                "received_fields": list(data.keys())
            }), 400

        # Validate work item IDs
        if not work_item_ids:
            return jsonify({
                "error": "Empty work item list",
                "message": "work_item_ids cannot be empty"
            }), 400

        # Convert all IDs to strings
        work_item_ids = [str(id) for id in work_item_ids]

        # Build URL and parameters
        ids_string = ",".join(work_item_ids)
        azure_url = f"https://dev.azure.com/dpwhotfsonline/_apis/wit/workitems"
        params = {
            "ids": ids_string,
            "fields": "System.WorkItemType,System.Title,System.Description,Microsoft.VSTS.Common.AcceptanceCriteria",
            "api-version": "7.1"
        }

        print(f"Azure URL: {azure_url}")  # Debug log
        print(f"Params: {params}")  # Debug log

        # Prepare authorization header
        if not token.startswith('eyJ'):  # Not a JWT token, so it's a PAT
            encoded_token = base64.b64encode(f":{token}".encode()).decode()
            auth_header = f"Basic {encoded_token}"
            print("Using Basic auth for PAT token")  # Debug log
        else:
            auth_header = f"Bearer {token}"
            print("Using Bearer auth")  # Debug log

        headers = {
            "Authorization": auth_header,
            "Content-Type": "application/json"
        }

        # Make request to Azure DevOps
        print(f"Making request to Azure DevOps...")  # Debug log
        response = requests.get(azure_url, headers=headers, params=params)

        print(f"Response status code: {response.status_code}")  # Debug log

        # Check if request was successful
        if response.status_code == 200:
            # Parse Azure DevOps response
            azure_data = response.json()
            print(f"Received {len(azure_data.get('value', []))} work items")  # Debug log

            # Process work items
            work_items = []
            for item in azure_data.get('value', []):
                fields = item.get('fields', {})

                workItemType = fields.get('System.WorkItemType', '')
                title = fields.get('System.Title', '')
                description_html = fields.get('System.Description', '')
                acceptance_criteria_html = fields.get('Microsoft.VSTS.Common.AcceptanceCriteria', '')

                # Convert HTML to plain text
                description_plain = clean_html_content(description_html)
                acceptance_criteria_plain = clean_html_content(acceptance_criteria_html)

                work_items.append({
                    "id": item.get('id'),
                    "workItemType" : workItemType,
                    "title": title,
                    "description": description_plain,
                    "acceptance_criteria": acceptance_criteria_plain
                })

            return jsonify({
                "success": True,
                "count": len(work_items),
                "work_items": work_items
            }), 200

        else:
            # Handle Azure DevOps API errors
            error_message = f"Azure DevOps API returned status code: {response.status_code}"
            error_detail = None

            try:
                error_detail = response.json()
                error_message = error_detail.get('message', error_message)
            except:
                error_message = response.text or error_message

            print(f"Azure API Error: {error_message}")  # Debug log

            return jsonify({
                "success": False,
                "error": "Azure DevOps API error",
                "message": error_message,
                "status_code": response.status_code,
                "detail": error_detail,
                "debug_info": {
                    "url": azure_url,
                    "params": params,
                    "work_item_ids": work_item_ids
                }
            }), response.status_code

    except requests.exceptions.ConnectionError as e:
        print(f"Connection error: {str(e)}")  # Debug log
        return jsonify({
            "success": False,
            "error": "Connection error",
            "message": "Failed to connect to Azure DevOps. Check your network connection.",
            "detail": str(e)
        }), 500

    except requests.exceptions.Timeout as e:
        print(f"Timeout error: {str(e)}")  # Debug log
        return jsonify({
            "success": False,
            "error": "Timeout error",
            "message": "Request to Azure DevOps timed out.",
            "detail": str(e)
        }), 500

    except requests.exceptions.RequestException as e:
        print(f"Request error: {str(e)}")  # Debug log
        return jsonify({
            "success": False,
            "error": "Network error",
            "message": str(e),
            "type": type(e).__name__
        }), 500

    except Exception as e:
        print(f"Unexpected error: {str(e)}")  # Debug log
        print(traceback.format_exc())  # Full stack trace

        return jsonify({
            "success": False,
            "error": "Internal server error",
            "message": str(e),
            "type": type(e).__name__,
            "traceback": traceback.format_exc() if app.debug else None
        }), 500


@app.route('/ai/api/workitemsDetails', methods=['POST'])
def get_work_items_details():
    """
    Get work items from Azure DevOps and merge acceptance criteria by type
    """
    try:
        # Get JSON payload
        data = request.get_json()
        print(f"Received data: {data}")  # Debug log

        # Validate payload
        if not data:
            return jsonify({
                "error": "No JSON payload provided",
                "debug": "Request body is empty or not valid JSON"
            }), 400

        # Check for token
        if 'token' not in data:
            return jsonify({
                "error": "Missing required field",
                "message": "'token' is required",
                "received_fields": list(data.keys())
            }), 400

        token = data['token']

        # Check for work_item_ids
        if 'work_item_ids' not in data:
            return jsonify({
                "error": "Missing required field",
                "message": "'work_item_ids' is required",
                "received_fields": list(data.keys())
            }), 400

        work_item_ids = data['work_item_ids']

        # Validate work_item_ids is a list
        if not isinstance(work_item_ids, list):
            return jsonify({
                "error": "Invalid format",
                "message": "'work_item_ids' must be a list",
                "received_type": type(work_item_ids).__name__
            }), 400

        # Validate work item IDs not empty
        if not work_item_ids:
            return jsonify({
                "error": "Empty work item list",
                "message": "work_item_ids cannot be empty"
            }), 400

        # Convert all IDs to strings
        work_item_ids = [str(id) for id in work_item_ids]

        # Build URL and parameters
        ids_string = ",".join(work_item_ids)
        azure_url = f"https://dev.azure.com/dpwhotfsonline/_apis/wit/workitems"
        params = {
            "ids": ids_string,
            "fields": "System.WorkItemType,System.Title,System.Description,Microsoft.VSTS.Common.AcceptanceCriteria",
            "api-version": "7.1"
        }

        print(f"Azure URL: {azure_url}")  # Debug log
        print(f"Params: {params}")  # Debug log

        # Prepare authorization header
        if not token.startswith('eyJ'):  # Not a JWT token, so it's a PAT
            encoded_token = base64.b64encode(f":{token}".encode()).decode()
            auth_header = f"Basic {encoded_token}"
            print("Using Basic auth for PAT token")  # Debug log
        else:
            auth_header = f"Bearer {token}"
            print("Using Bearer auth")  # Debug log

        headers = {
            "Authorization": auth_header,
            "Content-Type": "application/json"
        }

        # Make request to Azure DevOps
        print(f"Making request to Azure DevOps...")  # Debug log
        response = requests.get(azure_url, headers=headers, params=params)

        print(f"Response status code: {response.status_code}")  # Debug log

        # Check if request was successful
        if response.status_code == 200:
            # Parse Azure DevOps response
            azure_data = response.json()
            print(f"Received {len(azure_data.get('value', []))} work items")  # Debug log

            # Initialize containers for data
            epic_data = {
                "work_item_type": "Epic",
                "items": []
            }

            user_story_data = {
                "work_item_type": "User Story",
                "items": []
            }

            # Process work items
            for item in azure_data.get('value', []):
                fields = item.get('fields', {})

                workItemType = fields.get('System.WorkItemType', '')
                title = fields.get('System.Title', '')
                description_html = fields.get('System.Description', '')
                acceptance_criteria_html = fields.get('Microsoft.VSTS.Common.AcceptanceCriteria', '')

                # Convert HTML to plain text
                description_plain = clean_html_content(description_html)
                acceptance_criteria_plain = clean_html_content(acceptance_criteria_html)

                work_item_data = {
                    "id": item.get('id'),
                    "title": title,
                    "description": description_plain,
                    "acceptance_criteria": acceptance_criteria_plain
                }

                # Categorize by work item type
                if workItemType.lower() == 'epic':
                    epic_data["items"].append(work_item_data)
                elif workItemType.lower() == 'user story':
                    user_story_data["items"].append(work_item_data)

            # Create epic_merged_data
            epic_merged_data = ""
            if epic_data["items"]:
                epic_parts = []

                for item in epic_data["items"]:
                    epic_id = item['id']
                    epic_text_parts = []

                    # Add Epic with actual ID
                    epic_text_parts.append(f"Epic {epic_id}")

                    # Add description if exists
                    if item['description']:
                        epic_text_parts.append(f"the description is {item['description']}")

                    # Add acceptance criteria if exists
                    if item['acceptance_criteria'] and item['acceptance_criteria'].lower() != 'na':
                        epic_text_parts.append(f"and acceptance criteria is {item['acceptance_criteria']}")

                    # Join parts for this epic
                    if len(epic_text_parts) > 1:
                        epic_parts.append(" ".join(epic_text_parts))

                # Join all epics with comma and space
                epic_merged_data = ", ".join(epic_parts)

            # Create userstory_merged_data with new format
            userstory_merged_data = ""
            if user_story_data["items"]:
                story_parts = []

                for item in user_story_data["items"]:
                    story_id = item['id']
                    story_text_parts = []

                    # Add User Story with actual ID
                    story_text_parts.append(f"User Story {story_id}")

                    # Add description if exists
                    if item['description']:
                        story_text_parts.append(f"the description is {item['description']}")

                    # Add acceptance criteria if exists
                    if item['acceptance_criteria']:
                        story_text_parts.append(f"and acceptance criteria is {item['acceptance_criteria']}")

                    # Join parts for this story
                    if len(story_text_parts) > 1:
                        story_parts.append(" ".join(story_text_parts))

                # Join all stories with comma and space
                userstory_merged_data = ", ".join(story_parts)

            # Build response with ordered fields
            from collections import OrderedDict

            response_data = OrderedDict([
                ("success", True),
                ("total_items", len(azure_data.get('value', []))),
                ("epic_data", OrderedDict([
                    ("count", len(epic_data["items"])),
                    ("items", epic_data["items"]),
                    ("epic_merged_data", epic_merged_data)
                ])),
                ("user_story_data", OrderedDict([
                    ("count", len(user_story_data["items"])),
                    ("items", user_story_data["items"]),
                    ("userstory_merged_data", userstory_merged_data)
                ]))
            ])

            # Use json.dumps with Flask's response to maintain order
            return app.response_class(
                response=json.dumps(response_data),
                status=200,
                mimetype='application/json'
            )

        else:
            # Handle Azure DevOps API errors
            error_message = f"Azure DevOps API returned status code: {response.status_code}"
            error_detail = None

            try:
                error_detail = response.json()
                error_message = error_detail.get('message', error_message)
            except:
                error_message = response.text or error_message

            print(f"Azure API Error: {error_message}")  # Debug log

            return jsonify({
                "success": False,
                "error": "Azure DevOps API error",
                "message": error_message,
                "status_code": response.status_code,
                "detail": error_detail,
                "debug_info": {
                    "url": azure_url,
                    "params": params,
                    "work_item_ids": work_item_ids
                }
            }), response.status_code

    except requests.exceptions.ConnectionError as e:
        print(f"Connection error: {str(e)}")  # Debug log
        return jsonify({
            "success": False,
            "error": "Connection error",
            "message": "Failed to connect to Azure DevOps. Check your network connection.",
            "detail": str(e)
        }), 500

    except requests.exceptions.Timeout as e:
        print(f"Timeout error: {str(e)}")  # Debug log
        return jsonify({
            "success": False,
            "error": "Timeout error",
            "message": "Request to Azure DevOps timed out.",
            "detail": str(e)
        }), 500

    except requests.exceptions.RequestException as e:
        print(f"Request error: {str(e)}")  # Debug log
        return jsonify({
            "success": False,
            "error": "Network error",
            "message": str(e),
            "type": type(e).__name__
        }), 500

    except Exception as e:
        print(f"Unexpected error: {str(e)}")  # Debug log
        print(traceback.format_exc())  # Full stack trace

        return jsonify({
            "success": False,
            "error": "Internal server error",
            "message": str(e),
            "type": type(e).__name__,
            "traceback": traceback.format_exc() if app.debug else None
        }), 500


# Configure upload settings
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
ALLOWED_EXTENSIONS = {'.pdf', '.doc', '.docx', '.txt'}


class DocumentReader:
    """Class to handle reading different document formats"""

    @staticmethod
    def clean_text(text):
        """Clean and normalize text"""
        if not text:
            return ""

        # Remove line separators and dashes (like "-------")
        text = re.sub(r'[-─━]{3,}', ' ', text)

        # Remove other separator patterns (like "===", "***", "___")
        text = re.sub(r'[=*_]{3,}', ' ', text)

        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)

        # Remove special characters but keep basic punctuation
        text = re.sub(r'[^\w\s\.\,\!\?\:\;\-\(\)\"\']+', ' ', text)

        # Remove multiple spaces
        text = re.sub(r' +', ' ', text)

        # Remove newlines - replace with space
        text = re.sub(r'\n', ' ', text)

        # Replace quotes with single quotes for consistency
        text = re.sub(r'"', "'", text)

        # Clean up multiple spaces
        text = re.sub(r'\s+', ' ', text)

        # Clean up multiple periods or commas
        text = re.sub(r'\.+', '.', text)
        text = re.sub(r',+', ',', text)

        # Clean up spaces before/after punctuation
        text = re.sub(r'\s+([,.])', r'\1', text)
        text = re.sub(r'([,.])\s+', r'\1 ', text)

        return text.strip()

    @staticmethod
    def read_txt(file_path):
        """Read TXT file line by line with better cleaning"""
        lines = []
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    for line in file:
                        # Skip lines that are just separators
                        if re.match(r'^[-─━=*_\s]+$', line.strip()):
                            continue

                        cleaned_line = DocumentReader.clean_text(line.strip())
                        if cleaned_line:  # Only add non-empty lines
                            lines.append(cleaned_line)
                break
            except UnicodeDecodeError:
                continue
            except Exception as e:
                if encoding == encodings[-1]:  # Last encoding attempt
                    raise Exception(f"Failed to read file with any encoding: {str(e)}")

        return lines

    @staticmethod
    def read_pdf(file_path):
        """Read PDF file with improved text extraction"""
        lines = []
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()

                    if text:
                        # Split by newlines and clean
                        raw_lines = text.split('\n')
                        for line in raw_lines:
                            # Skip separator lines
                            if re.match(r'^[-─━=*_\s]+$', line.strip()):
                                continue

                            cleaned_line = DocumentReader.clean_text(line)
                            if cleaned_line:
                                lines.append(cleaned_line)

                        # Also split by sentences (periods followed by space and capital letter)
                        full_text = ' '.join(lines)
                        sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', full_text)

                        # Clear lines and add sentences
                        lines = []
                        for sentence in sentences:
                            cleaned_sentence = DocumentReader.clean_text(sentence)
                            if cleaned_sentence:
                                lines.append(cleaned_sentence)

        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")

        return lines

    @staticmethod
    def read_docx(file_path):
        """Read DOCX file with better paragraph handling"""
        lines = []
        try:
            doc = Document(file_path)

            # Read paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Skip separator lines
                    if re.match(r'^[-─━=*_\s]+$', paragraph.text.strip()):
                        continue

                    # Handle bullet points and numbered lists
                    text = paragraph.text

                    # Split by common delimiters
                    if '\n' in text:
                        sub_lines = text.split('\n')
                        for sub_line in sub_lines:
                            # Skip separator lines in sub-lines
                            if re.match(r'^[-─━=*_\s]+$', sub_line.strip()):
                                continue

                            cleaned_line = DocumentReader.clean_text(sub_line)
                            if cleaned_line:
                                lines.append(cleaned_line)
                    else:
                        cleaned_text = DocumentReader.clean_text(text)
                        if cleaned_text:
                            lines.append(cleaned_text)

            # Also read from tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = DocumentReader.clean_text(cell.text)
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        lines.append(" | ".join(row_text))

        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")

        return lines

    @staticmethod
    def read_doc(file_path):
        """Read DOC file (legacy format)"""
        try:
            # Try reading as DOCX first
            return DocumentReader.read_docx(file_path)
        except:
            # If you need to support .doc files, install python-docx2txt
            # pip install docx2txt
            try:
                import docx2txt
                text = docx2txt.process(file_path)
                lines = []
                for line in text.split('\n'):
                    # Skip separator lines
                    if re.match(r'^[-─━=*_\s]+$', line.strip()):
                        continue

                    cleaned_line = DocumentReader.clean_text(line)
                    if cleaned_line:
                        lines.append(cleaned_line)
                return lines
            except ImportError:
                raise Exception(
                    "Legacy .doc format not fully supported. Please install docx2txt or convert to .docx"
                )
            except Exception as e:
                raise Exception(f"Error reading .doc file: {str(e)}")


def allowed_file(filename):
    """Check if file extension is allowed"""
    return Path(filename).suffix.lower() in ALLOWED_EXTENSIONS


@app.route('/ai/upload', methods=['POST'])
def upload_document():
    """
    Upload and read document line by line

    Accepts: PDF, DOC, DOCX, TXT files
    Returns: List of lines from the document
    """

    # Check if file is in request
    if 'file' not in request.files:
        return jsonify({
            "status": "error",
            "message": "No file part in the request"
        }), 400

    file = request.files['file']

    # Check if file is selected
    if file.filename == '':
        return jsonify({
            "status": "error",
            "message": "No file selected"
        }), 400

    # Validate file extension
    if not allowed_file(file.filename):
        return jsonify({
            "status": "error",
            "message": f"File type not supported. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"
        }), 400

    # Get file extension
    file_extension = Path(file.filename).suffix.lower()

    # Create temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
        # Save uploaded file to temporary location
        file.save(tmp_file.name)
        tmp_file_path = tmp_file.name

    try:
        # Read document based on file type
        reader = DocumentReader()

        if file_extension == '.txt':
            lines = reader.read_txt(tmp_file_path)
        elif file_extension == '.pdf':
            lines = reader.read_pdf(tmp_file_path)
        elif file_extension == '.docx':
            lines = reader.read_docx(tmp_file_path)
        elif file_extension == '.doc':
            lines = reader.read_doc(tmp_file_path)
        else:
            raise Exception("Unsupported file type")

        # Additional cleaning - remove duplicates while preserving order
        seen = set()
        unique_lines = []
        for line in lines:
            if line not in seen:
                seen.add(line)
                unique_lines.append(line)

        # Prepare response
        response = {
            "filename": secure_filename(file.filename),
            "file_type": file_extension,
            "total_lines": len(unique_lines),
            "lines": unique_lines,
            "status": "success",
            "stats": {
                "original_lines": len(lines),
                "unique_lines": len(unique_lines),
                "duplicates_removed": len(lines) - len(unique_lines)
            }
        }

        return jsonify(response), 200

    except Exception as e:
        return jsonify({
            "status": "error",
            "message": f"Error processing document: {str(e)}"
        }), 500

    finally:
        # Clean up temporary file
        if os.path.exists(tmp_file_path):
            os.unlink(tmp_file_path)


@app.route('/ai/upload/raw', methods=['POST'])
def upload_document_raw():
    """
    Upload document and get raw text without line splitting
    Useful for documents where you want the full text content
    """

    # Check if file is in request
    if 'file' not in request.files:
        return jsonify({
            "status": "error",
            "message": "No file part in the request"
        }), 400

    file = request.files['file']

    # Check if file is selected
    if file.filename == '':
        return jsonify({
            "status": "error",
            "message": "No file selected"
        }), 400

    # Validate file extension
    if not allowed_file(file.filename):
        return jsonify({
            "status": "error",
            "message": f"File type not supported. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"
        }), 400

    # Get file extension
    file_extension = Path(file.filename).suffix.lower()

    # Create temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
        # Save uploaded file to temporary location
        file.save(tmp_file.name)
        tmp_file_path = tmp_file.name

    try:
        # Read document and get all lines
        reader = DocumentReader()

        if file_extension == '.txt':
            lines = reader.read_txt(tmp_file_path)
        elif file_extension == '.pdf':
            lines = reader.read_pdf(tmp_file_path)
        elif file_extension == '.docx':
            lines = reader.read_docx(tmp_file_path)
        elif file_extension == '.doc':
            lines = reader.read_doc(tmp_file_path)
        else:
            raise Exception("Unsupported file type")

        # Join all lines into full text
        full_text = '\n'.join(lines)

        # Prepare response
        response = {
            "filename": secure_filename(file.filename),
            "file_type": file_extension,
            "full_text": full_text,
            "character_count": len(full_text),
            "word_count": len(full_text.split()),
            "line_count": len(lines),
            "status": "success"
        }

        return jsonify(response), 200

    except Exception as e:
        return jsonify({
            "status": "error",
            "message": f"Error processing document: {str(e)}"
        }), 500

    finally:
        # Clean up temporary file
        if os.path.exists(tmp_file_path):
            os.unlink(tmp_file_path)



class ChatBoxLLM:
    def __init__(self):
        """Initialize the LLM and conversation history"""
        self.llm = ChatGoogleGenerativeAI(
            model="models/gemini-2.0-flash",
            api_key=os.getenv('GOOGLE_API_KEY')
        )
        self.conversation_history = []

    def get_response(self, prompt: str) -> str:
        """
        Process a prompt and return response while maintaining conversation context

        Args:
            prompt: The user's input message

        Returns:
            The LLM's response as a string
        """
        # Add user message to history
        self.conversation_history.append(HumanMessage(content=prompt))

        # Get response from LLM with full conversation history
        response = self.llm.invoke(self.conversation_history)

        # Add AI response to history
        self.conversation_history.append(AIMessage(content=response.content))

        return response.content.strip()

    def clear_history(self):
        """Clear the conversation history"""
        self.conversation_history = []

    def get_history(self):
        """Get the current conversation history"""
        return [
            {
                "role": "human" if isinstance(msg, HumanMessage) else "assistant",
                "content": msg.content
            }
            for msg in self.conversation_history
        ]


# Initialize the chat bot instance globally
chat_bot = ChatBoxLLM()


@app.route('/ai/chat-box-llm', methods=['POST'])
def chat_box_llm():
    try:
        data = request.json or {}
        prompt = data.get("prompt", "").strip()

        if not prompt:
            return Response("Prompt is required", status=400, mimetype="text/plain")

        print(f"📥 Received Prompt: {prompt}")

        # Use the chat bot instance to get response
        result = chat_bot.get_response(prompt)

        print(f"💬 Gemini Response: {result}")

        return Response(result, mimetype="text/plain", status=200)

    except Exception as e:
        print(f"❌ Error: {str(e)}")
        return Response(f"Error: {str(e)}", status=500, mimetype="text/plain")


# Optional: Add endpoint to clear conversation history
@app.route('/ai/chat-box-llm/clear', methods=['POST'])
def clear_chat_history():
    try:
        chat_bot.clear_history()
        return Response("Conversation history cleared", status=200, mimetype="text/plain")
    except Exception as e:
        return Response(f"Error: {str(e)}", status=500, mimetype="text/plain")


# Optional: Add endpoint to get conversation history
@app.route('/ai/chat-box-llm/history', methods=['GET'])
def get_chat_history():
    try:
        history = chat_bot.get_history()
        return jsonify(history)
    except Exception as e:
        return Response(f"Error: {str(e)}", status=500, mimetype="text/plain")

@app.route('/ai/health', methods=['GET'])
def health_check():
    try:
        if hasattr(chat_bot, 'llm'):
            return jsonify({"status": "healthy"}), 200
        else:
            return jsonify({"status": "unhealthy", "issue": "chat_bot_not_initialized"}), 503
    except:
        return jsonify({"status": "unhealthy"}), 503


@app.errorhandler(413)
def too_large(e):
    return jsonify({
        "status": "error",
        "message": "File is too large. Maximum size is 16MB"
    }), 413


if __name__ == '__main__':
    print("🚀 Enhanced New Relic API Discovery Suite with Integrated Parameter Extraction")
    print("=" * 80)
    print("✅ Landing page available at http://localhost:5000")
    print("✅ LLM API Discovery available at http://localhost:5000/llmapidiscovery")
    print("\n✅ Available endpoints:")
    print("  POST /get-accounts           - Get available accounts")
    print("  POST /get-apps               - Get available apps for an account")
    print("  POST /test-connection        - Test connection with exact app name validation")
    print("  POST /analyze-data           - Discover all APIs from last 30 days")
    print("  POST /get-endpoint-details   - Get comprehensive endpoint details with parameters & cURL")
    print("  POST /export-data            - Export API analysis data")
    print("  POST /validate-app           - Validate selected app name")
    print("\n🆕 ENHANCED FEATURES:")
    print("  • Integrated parameter extraction in /get-endpoint-details")
    print("  • Complete Path, Query, Header, and Body parameter detection")
    print("  • Automatic cURL command generation")
    print("  • Nested object and array parameter handling")
    print("  • Enhanced request/response payload schemas")
    print("  • Single endpoint for all endpoint information")
    print("=" * 80)
    app.run(debug=True, host='0.0.0.0', port=5000)